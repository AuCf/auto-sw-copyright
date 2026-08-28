"""
CPCC-Compliant Docx Generation Engine
Produces standard 60-page Source Code document (50 lines/page) and User Manual document with embedded UI screenshots.
"""

import io
import re
from typing import List, Optional, Dict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_run_font(run, font_name="Microsoft YaHei", size_pt=10.5, bold=False, color_rgb=None):
    """Safely apply font with EastAsia Chinese character binding."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb

def add_xml_page_number(run):
    """Insert dynamic Word PAGE field into a text run."""
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    run._r.append(fldSimple)

def add_xml_total_pages(run):
    """Insert dynamic Word NUMPAGES field into a text run."""
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'NUMPAGES')
    run._r.append(fldSimple)

def set_cell_background(cell, fill_hex="F3F4F6"):
    """Set table cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def sanitize_text(text: str) -> str:
    """Remove control characters that cause Word XML parsing errors or garbled text."""
    if not text:
        return ""
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
    return text

def insert_image_into_doc(doc: Document, image_bytes: bytes, caption_text: str):
    """Safely insert a center-aligned image with caption into the docx."""
    try:
        p_img = doc.add_paragraph()
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(io.BytesIO(image_bytes), width=Inches(5.6))

        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(2)
        p_cap.paragraph_format.space_after = Pt(12)
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run(caption_text)
        set_run_font(r_cap, font_name="SimSun", size_pt=9.0, bold=True, color_rgb=RGBColor(100, 116, 139))
    except Exception as e:
        print(f"Failed to insert image: {e}")


def generate_code_docx(
    code_lines: List[str],
    software_name: str,
    version: str = "V1.0",
    lines_per_page: int = 50,
    target_pages: int = 60
) -> io.BytesIO:
    """
    Generate standard 60-page CPCC source code Word document.
    Strictly enforce 50 lines per page, with standardized headers and footers.
    """
    doc = Document()
    
    # 1. Configure page margins (2.54cm / 1 inch)
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    # 2. Extract code lines (First 30 pages + Last 30 pages if total > 60 pages)
    total_required = target_pages * lines_per_page
    cleaned_lines = [sanitize_text(line.rstrip()) for line in code_lines if line is not None]
    
    if len(cleaned_lines) > total_required:
        half_needed = (target_pages // 2) * lines_per_page
        selected_lines = cleaned_lines[:half_needed] + cleaned_lines[-half_needed:]
    else:
        selected_lines = cleaned_lines

    # 3. Setup Header & Footer
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = ""
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Left side: Software Full Name + Version + Code Document
    run_left = header_para.add_run(f"《{software_name}》 {version} 源程序代码文档\t\t")
    set_run_font(run_left, font_name="SimSun", size_pt=8.5, color_rgb=RGBColor(100, 100, 100))
    
    # Right side: Page Number
    run_page_prefix = header_para.add_run("第 ")
    set_run_font(run_page_prefix, font_name="SimSun", size_pt=8.5, color_rgb=RGBColor(100, 100, 100))
    
    run_page = header_para.add_run()
    set_run_font(run_page, font_name="Times New Roman", size_pt=8.5, color_rgb=RGBColor(100, 100, 100))
    add_xml_page_number(run_page)
    
    run_page_mid = header_para.add_run(" 页 / 共 ")
    set_run_font(run_page_mid, font_name="SimSun", size_pt=8.5, color_rgb=RGBColor(100, 100, 100))
    
    run_total = header_para.add_run(f"{target_pages}")
    set_run_font(run_total, font_name="Times New Roman", size_pt=8.5, color_rgb=RGBColor(100, 100, 100))
    
    run_page_suffix = header_para.add_run(" 页")
    set_run_font(run_page_suffix, font_name="SimSun", size_pt=8.5, color_rgb=RGBColor(100, 100, 100))

    # 4. Render code paragraphs with exact line counts
    current_page = 1
    line_in_current_page = 0
    total_rendered_lines = 0

    for line_text in selected_lines:
        if current_page > target_pages:
            break
            
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(11.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        total_rendered_lines += 1
        line_num_str = f"{total_rendered_lines:4d}  "
        
        r_num = p.add_run(line_num_str)
        set_run_font(r_num, font_name="Consolas", size_pt=8.5, color_rgb=RGBColor(160, 160, 160))
        
        r_code = p.add_run(line_text if line_text else " ")
        set_run_font(r_code, font_name="Consolas", size_pt=8.5, color_rgb=RGBColor(30, 41, 59))
        
        line_in_current_page += 1
        
        if line_in_current_page == lines_per_page and current_page < target_pages:
            doc.add_page_break()
            current_page += 1
            line_in_current_page = 0

    while current_page <= target_pages and total_rendered_lines < total_required:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(11.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        total_rendered_lines += 1
        r_num = p.add_run(f"{total_rendered_lines:4d}  ")
        set_run_font(r_num, font_name="Consolas", size_pt=8.5, color_rgb=RGBColor(160, 160, 160))
        
        r_code = p.add_run("// End of module block")
        set_run_font(r_code, font_name="Consolas", size_pt=8.5, color_rgb=RGBColor(140, 140, 140))
        
        line_in_current_page += 1
        if line_in_current_page == lines_per_page and current_page < target_pages:
            doc.add_page_break()
            current_page += 1
            line_in_current_page = 0

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def generate_manual_docx(
    markdown_content: str,
    software_name: str,
    version: str = "V1.0",
    images: Optional[Dict[str, bytes]] = None
) -> io.BytesIO:
    """
    Convert User Manual Markdown content into a rich Word document with embedded UI screenshots.
    """
    doc = Document()
    images = images or {}
    
    # 1. Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # 2. Cover / Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(40)
    title_p.paragraph_format.space_after = Pt(15)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = title_p.add_run(f"《{sanitize_text(software_name)}》")
    set_run_font(r_title, font_name="Microsoft YaHei", size_pt=22, bold=True, color_rgb=RGBColor(15, 23, 42))

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(50)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub_p.add_run(f"用户操作手册与系统设计说明书 ({version})")
    set_run_font(r_sub, font_name="Microsoft YaHei", size_pt=15, bold=False, color_rgb=RGBColor(71, 85, 105))

    # 3. Setup Header/Footer
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_head = header_p.add_run(f"《{sanitize_text(software_name)}》 {version} 用户手册\t\t第 ")
    set_run_font(r_head, font_name="SimSun", size_pt=9, color_rgb=RGBColor(120, 120, 120))
    
    r_page = header_p.add_run()
    set_run_font(r_page, font_name="Times New Roman", size_pt=9, color_rgb=RGBColor(120, 120, 120))
    add_xml_page_number(r_page)
    
    r_head_end = header_p.add_run(" 页")
    set_run_font(r_head_end, font_name="SimSun", size_pt=9, color_rgb=RGBColor(120, 120, 120))

    # 4. Parse Markdown Lines & Inject UI Images
    lines = markdown_content.splitlines()
    in_table = False
    table_rows = []
    
    inserted_flags = {
        "login": False,
        "dashboard": False,
        "module_1": False,
        "module_2": False,
        "module_3": False,
        "module_4": False,
        "config_modal": False,
    }

    current_module_idx = 1

    for line in lines:
        stripped = sanitize_text(line.strip())
        if not stripped:
            if in_table and table_rows:
                _render_table(doc, table_rows)
                in_table = False
                table_rows = []
            continue

        # Markdown Table Detection
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            if not re.match(r"^\|(\s*:?-+:?\s*\|)+$", stripped):
                cells = [c.strip() for c in stripped.split("|")[1:-1]]
                table_rows.append(cells)
            continue
        elif in_table and table_rows:
            _render_table(doc, table_rows)
            in_table = False
            table_rows = []

        # Heading 1 (#)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(stripped[2:].strip())
            set_run_font(r, font_name="Microsoft YaHei", size_pt=16, bold=True, color_rgb=RGBColor(15, 23, 42))

        # Heading 2 (##)
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(stripped[3:].strip())
            set_run_font(r, font_name="Microsoft YaHei", size_pt=13, bold=True, color_rgb=RGBColor(30, 58, 138))

            # Smart Image Injection for Chapters
            h_text = stripped[3:].strip()
            if "系统概述" in h_text or "总体架构" in h_text:
                if not inserted_flags["dashboard"] and "dashboard" in images:
                    insert_image_into_doc(doc, images["dashboard"], f"图 1-1 《{software_name}》 总体运行监控与大屏看板示意图")
                    inserted_flags["dashboard"] = True
            elif "系统登录" in h_text or "身份认证" in h_text:
                if not inserted_flags["login"] and "login" in images:
                    insert_image_into_doc(doc, images["login"], "图 3-1 系统统一身份认证与安全登录界面示意图")
                    inserted_flags["login"] = True

        # Heading 3 (###)
        elif stripped.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(stripped[4:].strip())
            set_run_font(r, font_name="Microsoft YaHei", size_pt=11, bold=True, color_rgb=RGBColor(51, 65, 85))

            h3_text = stripped[4:].strip()
            # Module UI screen injection
            if any(k in h3_text for k in ["操作界面", "界面示意", "功能操作", "数据列表"]):
                mod_key = f"module_{current_module_idx}"
                if mod_key in images and not inserted_flags.get(mod_key, False):
                    insert_image_into_doc(doc, images[mod_key], f"图 {current_module_idx + 3}-1 业务数据列表与操作工作台示意图")
                    inserted_flags[mod_key] = True
                    current_module_idx += 1
                elif "config_modal" in images and not inserted_flags["config_modal"]:
                    insert_image_into_doc(doc, images["config_modal"], "图 4-2 核心业务参数与告警阈值配置弹窗示意图")
                    inserted_flags["config_modal"] = True

        # Bullet list (- or *)
        elif stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _render_rich_text(p, stripped[2:].strip())

        # Numbered list (1. 2.)
        elif re.match(r"^\d+\.\s+", stripped):
            match = re.match(r"^\d+\.\s+", stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _render_rich_text(p, stripped[match.end():].strip())

        # Mermaid code block placeholder
        elif stripped.startswith("```mermaid"):
            if not inserted_flags["dashboard"] and "dashboard" in images:
                insert_image_into_doc(doc, images["dashboard"], "图 1-1 系统总体架构与数据流转拓扑图")
                inserted_flags["dashboard"] = True
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run("【系统总体架构与数据流转拓扑图】")
                set_run_font(r, font_name="Microsoft YaHei", size_pt=9.5, bold=True, color_rgb=RGBColor(59, 130, 246))

        elif stripped.startswith("```"):
            continue

        # Normal Paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = Pt(15)
            _render_rich_text(p, stripped)

    if in_table and table_rows:
        _render_table(doc, table_rows)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _render_rich_text(paragraph, text: str):
    """Parse inline markdown elements (bold **text**, code `text`)."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            r = paragraph.add_run(part[2:-2])
            set_run_font(r, font_name="SimSun", size_pt=10.5, bold=True)
        else:
            sub_parts = re.split(r"(`.*?`)", part)
            for sub in sub_parts:
                if sub.startswith("`") and sub.endswith("`") and len(sub) >= 2:
                    r = paragraph.add_run(sub[1:-1])
                    set_run_font(r, font_name="Consolas", size_pt=9.5, color_rgb=RGBColor(220, 38, 38))
                else:
                    r = paragraph.add_run(sub)
                    set_run_font(r, font_name="SimSun", size_pt=10.5, color_rgb=RGBColor(30, 41, 59))


def _render_table(doc: Document, rows: List[List[str]]):
    """Render a clean styled table in Word with font bindings."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.autofit = True
    
    for row_idx, row in enumerate(rows):
        for col_idx in range(num_cols):
            cell_text = row[col_idx] if col_idx < len(row) else ""
            cell = table.cell(row_idx, col_idx)
            cell.text = cell_text
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                if row_idx == 0:
                    set_run_font(run, font_name="SimSun", size_pt=9.5, bold=True, color_rgb=RGBColor(15, 23, 42))
                else:
                    set_run_font(run, font_name="SimSun", size_pt=9.5, color_rgb=RGBColor(51, 65, 85))
            
            if row_idx == 0:
                set_cell_background(cell, "E2E8F0")
            elif row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
                
    doc.add_paragraph()  # spacing after table
